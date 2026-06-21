from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "mode-emploi-plateforme-mykori.docx"
CAPTURES = ROOT / "docs" / "mode-emploi-captures"
LOGO = ROOT / "public" / "assets" / "images" / "kori" / "kori-logo.png"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
GOLD = "F2B400"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4CC"
LIGHT_GREEN = "E8F7EF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill=LIGHT_BLUE):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(INK)


def add_title(doc, title, subtitle):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Inches(1.25))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    widths = [Inches(1.6), Inches(4.7)]
    labels = [
        ("Plateforme", "myKORI - KORI Asset Management"),
        ("Profil documente", "Espace Asset Manager / KAM"),
        ("Date", "21/06/2026"),
        ("Source des captures", "Environnement local mykori.test"),
    ]
    for row, (label, value) in zip(meta.rows, labels):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_text(row.cells[0], label, bold=True, color="0B2545")
        set_cell_text(row.cells[1], value)
    style_table(meta, LIGHT_GOLD)

    note(
        doc,
        "Ce guide explique l'utilisation operationnelle de la plateforme: connexion, clients, placements, "
        "valorisation, releves, liquidite et preuves de paiement. Les ecrans peuvent varier selon le role, "
        "les droits et les donnees visibles dans l'environnement.",
        fill=LIGHT_GREEN,
    )
    doc.add_page_break()


def note(doc, text, fill=LIGHT_GRAY, title=None):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if title:
        r_title = p.add_run(title + " ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor.from_string("0B2545")
        r_title.font.size = Pt(9.5)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(step)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)


def add_capture(doc, filename, caption):
    path = CAPTURES / filename
    if not path.exists():
        note(doc, f"Capture manquante: {filename}", fill="FEE2E2")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_section(doc, title, body=None):
    doc.add_heading(title, level=1)
    if body:
        p = doc.add_paragraph(body)
        p.paragraph_format.space_after = Pt(8)


def add_subsection(doc, title, body=None):
    doc.add_heading(title, level=2)
    if body:
        p = doc.add_paragraph(body)
        p.paragraph_format.space_after = Pt(6)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Mode d'emploi myKORI - KORI Asset Management"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Document interne - version du 21/06/2026"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    add_title(
        doc,
        "Mode d'emploi complet de la plateforme myKORI",
        "Guide utilisateur avec captures d'ecran et regles de gestion PMG/FCP",
    )

    doc.add_heading("Sommaire", level=1)
    add_bullets(
        doc,
        [
            "Acces, connexion et principes de securite",
            "Navigation Asset Manager et tableau de bord",
            "Gestion des clients et creation d'un compte client",
            "Gestion des placements PMG et FCP",
            "Valorisation, liquidite, paiements et justificatifs",
            "Generation, controle et envoi des releves clients",
            "Controles mensuels et depannage courant",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "1. Acces et connexion",
        "La plateforme est accessible depuis l'adresse fournie par l'administrateur. En local, les captures de ce guide utilisent mykori.test. Chaque utilisateur doit se connecter avec son identifiant et son mot de passe, puis respecter les controles de securite actives sur son profil.",
    )
    add_capture(doc, "01-connexion.png", "Ecran de connexion de la plateforme myKORI.")
    add_steps(
        doc,
        [
            "Ouvrir l'adresse de la plateforme dans le navigateur.",
            "Saisir l'email et le mot de passe du compte autorise.",
            "Valider la connexion et suivre l'eventuelle etape de verification.",
            "Se deconnecter en fin de session, surtout sur un poste partage.",
        ],
    )
    note(
        doc,
        "Un utilisateur ne doit jamais partager ses identifiants. Toute operation sensible, notamment un paiement ou une regularisation de liquidite, doit rester rattachee au compte de la personne qui l'a enregistree.",
        fill=LIGHT_GOLD,
        title="Securite.",
    )

    add_section(
        doc,
        "2. Roles et responsabilites",
        "Les menus visibles dependent du role connecte. Ce guide couvre principalement l'espace Asset Manager, mais les regles ci-dessous servent de reference pour les controles croises.",
    )
    roles = doc.add_table(rows=1, cols=3)
    roles.style = "Table Grid"
    hdr = roles.rows[0].cells
    for i, text in enumerate(["Role", "Responsabilite principale", "Points de controle"]):
        set_cell_text(hdr[i], text, bold=True)
    rows = [
        ("Asset Manager / KAM", "Suit les clients, placements, portefeuilles et releves.", "Verifier les montants, echeances, liquidites et releves avant envoi."),
        ("Backoffice", "Enregistre ou valide les operations de gestion.", "Controler les dates, pieces justificatives et references de paiement."),
        ("Client", "Consulte son portefeuille, ses releves et ses liquidites.", "Les informations affichees doivent etre coherentes avec les operations validees."),
        ("Administrateur", "Gere les acces, profils et parametrages.", "Limiter les droits au besoin reel et auditer les comptes actifs."),
    ]
    for role, responsibility, control in rows:
        cells = roles.add_row().cells
        set_cell_text(cells[0], role, bold=True, size=9)
        set_cell_text(cells[1], responsibility, size=9)
        set_cell_text(cells[2], control, size=9)
    style_table(roles)

    add_subsection(
        doc,
        "Cartographie des espaces de la plateforme",
        "Les routes du projet montrent plusieurs espaces fonctionnels. L'acces a chaque espace depend des permissions attribuees a l'utilisateur connecte.",
    )
    modules = doc.add_table(rows=1, cols=4)
    modules.style = "Table Grid"
    for i, text in enumerate(["Espace", "Acces principal", "Utilisation", "Utilisateurs concernes"]):
        set_cell_text(modules.rows[0].cells[i], text, bold=True)
    for row in [
        ("Portail client", "/dashboard, /my-products, /my-statements", "Consultation portefeuille, produits, historiques et releves.", "Clients"),
        ("Asset Manager", "/asset-manager", "Gestion clients, placements, transactions, VL et releves.", "KAM / Asset managers"),
        ("Compliance", "/compliance", "Controle clients, audit portefeuille, historique VL et releves.", "Conformite"),
        ("Backoffice", "/backoffice", "Suivi et validation des transactions.", "Backoffice / validateurs"),
        ("Direction Generale", "/dg", "Vue de pilotage et historique des campagnes de releves.", "DG"),
        ("Admin Frontend", "/admin-front", "Gestion utilisateurs, menus et logs d'activite.", "Administrateurs"),
        ("CRM", "/crm", "Prospects, statuts commerciaux et conversion en clients.", "Equipe commerciale"),
        ("Admin technique", "/admin", "Administration Voyager/back-office technique.", "Administrateurs techniques"),
    ]:
        cells = modules.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=8.3)
    style_table(modules, LIGHT_BLUE)

    add_section(
        doc,
        "3. Navigation Asset Manager",
        "La navigation laterale permet d'acceder aux principales fonctions: tableau de bord, clients, creation client, valeurs liquidatives et deconnexion. Les boutons en haut de certaines pages donnent acces aux transactions, aux historiques et aux releves.",
    )
    add_capture(doc, "02-dashboard-asset-manager.png", "Tableau de bord Asset Manager.")
    add_bullets(
        doc,
        [
            "Tableau de bord: synthese des encours, gains, clients actifs et alertes.",
            "Clients: consultation de la base client, filtrage par type de produit et acces aux details.",
            "Creer un client: creation d'un nouveau compte et des informations d'identification.",
            "Valeurs Liquidatives: saisie et controle des VL utilisees pour les produits FCP.",
            "Transactions et historiques: suivi des operations d'un client ou d'un produit.",
        ],
    )

    add_section(
        doc,
        "4. Tableau de bord et indicateurs",
        "Le tableau de bord agrege les donnees issues des clients, placements et mouvements financiers. Les indicateurs ne doivent pas etre lus comme une preuve comptable isolee: ils doivent etre reconciles avec les operations detaillees.",
    )
    indicators = doc.add_table(rows=1, cols=3)
    indicators.style = "Table Grid"
    for i, text in enumerate(["Indicateur", "Interpretation", "Controle attendu"]):
        set_cell_text(indicators.rows[0].cells[i], text, bold=True)
    for row in [
        ("Capital investi", "Somme des montants places et encore pertinents selon la vue.", "Comparer aux placements actifs et aux remboursements."),
        ("Gains FCP", "Gains calcules par difference entre valeur actuelle et cout de revient.", "Verifier VL et nombre de parts avec toutes les decimales."),
        ("Gains actifs PMG", "Interets calcules sur les PMG actifs ou a declarer.", "Verifier date de debut, date de fin, taux et mode de gestion des interets."),
        ("Clients actifs", "Clients avec portefeuille ou produits actifs selon le filtre.", "Ne pas confondre avec la liste des releves a envoyer."),
        ("Clients inactifs", "Clients sans portefeuille actif dans le perimetre affiche.", "Verifier qu'aucun placement en cours ou liquidite residuelle n'existe."),
    ]:
        cells = indicators.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=9)
    style_table(indicators)

    add_section(
        doc,
        "5. Gestion des clients",
        "La liste clients sert a rechercher un client, consulter ses produits, verifier son portefeuille global et ouvrir sa fiche detaillee. Les onglets permettent de distinguer tous les clients, les clients FCP, les clients PMG, les echeances du mois et les anniversaires PMG.",
    )
    add_capture(doc, "03-liste-clients.png", "Liste generale des clients et indicateurs de synthese.")
    add_capture(doc, "04-liste-clients-pmg.png", "Filtrage de la liste sur les clients PMG.")
    add_bullets(
        doc,
        [
            "Cliquer sur Details pour ouvrir la fiche d'un client.",
            "Utiliser les onglets pour verifier les populations PMG, FCP, echeances et anniversaires.",
            "Un client peut apparaitre avec un portefeuille a zero s'il existe dans la base mais sans placement actif dans le filtre courant.",
            "Les clients selectionnes pour les releves peuvent differer du nombre de clients actifs, car la logique d'envoi depend aussi de la periode de valorisation et des produits a declarer.",
        ],
    )

    add_section(
        doc,
        "6. Creation d'un client",
        "La creation client doit etre complete, car les donnees sont ensuite utilisees pour les releves, les emails et l'identification du portefeuille.",
    )
    add_capture(doc, "10-creation-client.png", "Formulaire de creation d'un nouveau client.")
    add_steps(
        doc,
        [
            "Ouvrir Creer un client depuis le menu lateral.",
            "Renseigner les informations d'identite ou de raison sociale.",
            "Saisir une adresse email valide pour les notifications et releves.",
            "Verifier le role, le type de client et les informations de contact.",
            "Enregistrer puis controler que le client apparait dans la liste.",
        ],
    )

    add_section(
        doc,
        "7. Fiche client et portefeuille",
        "La fiche client consolide les produits PMG/FCP, les interets generes, le portefeuille global, les liquidites et les operations accessibles pour ce client.",
    )
    add_capture(doc, "05-detail-client-africa-bright.png", "Exemple de fiche client: AFRICA BRIGHT.")
    add_bullets(
        doc,
        [
            "Portefeuille PMG: capital et interets lies aux mandats de gestion.",
            "Portefeuille FCP: valeur calculee a partir du nombre de parts et de la VL applicable.",
            "Interets generes: total des interets produits selon les regles de chaque placement.",
            "Liquidite: montants sortis du placement mais pas encore payes au client.",
            "Historique de transactions: preuve principale des mouvements enregistres.",
        ],
    )

    add_section(
        doc,
        "8. Portail client",
        "Le portail client donne au client une vision de ses produits, de ses historiques et de ses releves. Les informations visibles doivent etre strictement coherentes avec les operations validees et les liquidites non encore payees.",
    )
    add_bullets(
        doc,
        [
            "Dashboard client: synthese du portefeuille et acces aux rubriques personnelles.",
            "Mes produits: detail des produits detenus, gains et evolution.",
            "Mon historique: liste des transactions et telechargement possible d'un historique PDF.",
            "Mes releves: releves mensuels disponibles par annee, mois et type de produit.",
            "Profil: informations du compte et changement de mot de passe si autorise.",
        ],
    )
    note(
        doc,
        "Une liquidite doit etre visible pour le client tant qu'elle n'est pas payee. Apres paiement, l'ecriture de paiement et son justificatif deviennent la preuve de sortie effective.",
        fill=LIGHT_GREEN,
        title="Regle portail client.",
    )

    add_section(
        doc,
        "9. Operations et transactions",
        "La page de gestion des transactions permet d'ajouter, verifier et historiser les mouvements d'un client: placements, rachats, remboursements, paiements d'interets et paiements de liquidite.",
    )
    add_capture(doc, "06-operations-client-africa-bright.png", "Gestion des operations d'un client.")
    add_bullets(
        doc,
        [
            "Chaque operation doit avoir une date, un montant et un type coherent.",
            "Les operations PMG doivent respecter le mode de gestion des interets du mandat.",
            "Les paiements au client doivent porter une preuve: date de paiement, methode, reference et justificatif.",
            "L'historique doit permettre de reconstituer le solde, les liquidites et les paiements deja effectues.",
        ],
    )

    add_section(
        doc,
        "10. Regles PMG: capitalisation, distribution et liquidite",
        "Un PMG doit etre interprete a partir de son montant place, de son taux, de ses dates, et surtout de sa modalite de gestion des interets.",
    )
    pmg = doc.add_table(rows=1, cols=3)
    pmg.style = "Table Grid"
    for i, text in enumerate(["Modalite", "Traitement attendu", "Effet sur le client"]):
        set_cell_text(pmg.rows[0].cells[i], text, bold=True)
    for row in [
        ("A l'echeance du placement", "Les interets courent jusqu'a la date de fin du mandat. A l'echeance, capital + interets doivent basculer en liquidite si le paiement n'est pas encore execute.", "Le client voit une liquidite a payer, puis un paiement prouve solde cette liquidite."),
        ("Annuellement a date anniversaire", "A chaque anniversaire, les interets acquis doivent etre mis en liquidite automatiquement si le backoffice n'a pas saisi la transaction a temps.", "Les interets sont visibles dans l'espace client et sur les releves."),
        ("Capitalisation jusqu'a l'echeance", "Les interets restent dans le placement jusqu'a la fin du mandat et s'ajoutent a la valeur finale selon la regle de calcul retenue.", "Pas de paiement intermediaire; sortie a l'echeance."),
        ("Interets precomptes", "Les interets sont constates au depart selon le montage prevu; ils doivent etre traces comme mouvement distinct.", "Le client voit l'effet financier des le demarrage."),
        ("Mensuel exceptionnel", "Les interets du mois doivent etre mis en liquidite automatiquement a la date mensuelle prevue si aucune saisie manuelle n'est faite.", "Le client voit la liquidite mensuelle et le releve doit la reprendre."),
    ]:
        cells = pmg.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=8.5)
    style_table(pmg, LIGHT_GOLD)
    note(
        doc,
        "Principe de preuve: une liquidite n'est pas un paiement. Elle prouve qu'un montant est exigible ou disponible. Le paiement est prouve uniquement par une ecriture de paiement avec date, methode, reference et justificatif.",
        fill=LIGHT_GOLD,
        title="Important.",
    )

    add_section(
        doc,
        "11. Valorisation mensuelle",
        "La date de valorisation mensuelle affichee a la date N doit etre le dernier jour du mois. Exemple: pour les releves de juin, la valorisation de reference du mois precedent est le 30/06 si le mois est juin, ou le 31/05 pour une campagne d'envoi debut juin.",
    )
    add_bullets(
        doc,
        [
            "Par defaut, la valorisation est arretee au dernier jour du mois.",
            "Si un PMG expire pendant le mois et qu'il est seul dans le releve concerne, sa valorisation s'arrete a la date de fin de mandat.",
            "Si le PMG expire au cours du mois mais figure avec d'autres produits du client, la vue portefeuille globale conserve le dernier jour du mois pour l'ensemble.",
            "Entre le 1er et le 5 du mois suivant, les produits a relever doivent rester visibles pour permettre l'envoi des releves meme si une expiration est intervenue.",
            "Un placement effectue le dernier jour du mois doit etre pris en compte si sa date d'effet est dans la periode valorisee.",
        ],
    )

    add_section(
        doc,
        "12. Regles FCP et valeurs liquidatives",
        "Pour les FCP, le calcul doit utiliser toutes les decimales disponibles de la valeur liquidative et du nombre de parts. L'arrondi a deux decimales ne sert qu'a l'affichage client.",
    )
    add_capture(doc, "09-valeurs-liquidatives.png", "Ecran de suivi des valeurs liquidatives.")
    add_bullets(
        doc,
        [
            "Calcul interne: nombre de parts exact x VL exacte, sans tronquer les decimales.",
            "Affichage client: montant presente avec deux chiffres apres la virgule.",
            "Controle: comparer les calculs au centime avant envoi des releves.",
            "Anomalie typique: une VL arrondie trop tot peut creer un ecart visible sur les gros portefeuilles.",
        ],
    )

    add_section(
        doc,
        "13. Releves clients PMG et FCP",
        "Les releves sont generes a partir des portefeuilles, gains et liquidites visibles dans la periode de campagne. La liste d'envoi ne doit pas etre confondue avec le simple nombre de clients actifs.",
    )
    add_capture(doc, "07-releves-pmg.png", "Liste de generation et d'envoi des releves PMG.")
    add_capture(doc, "08-releves-fcp.png", "Liste de generation et d'envoi des releves FCP.")
    add_steps(
        doc,
        [
            "Ouvrir la liste des releves PMG ou FCP.",
            "Verifier les clients coches et les montants affiches.",
            "Previsualiser les releves sensibles ou ceux avec montant a zero.",
            "Corriger les anomalies de donnees avant envoi.",
            "Generer et envoyer les releves uniquement apres validation des montants.",
        ],
    )
    note(
        doc,
        "Les valeurs nulles dans la liste d'envoi indiquent souvent un client inclus dans le perimetre mais sans portefeuille valorisable dans la periode courante, ou un filtre qui retient le client pour historique/releve sans montant actif. Ces lignes doivent etre verifiees avant envoi.",
        fill=LIGHT_GOLD,
        title="Controle releves.",
    )

    add_section(
        doc,
        "14. Paiements, liquidites et preuve",
        "La plateforme doit distinguer trois etats: montant investi, liquidite disponible, paiement effectue. Le passage en liquidite ne signifie pas que le client a ete paye.",
    )
    proof = doc.add_table(rows=1, cols=4)
    proof.style = "Table Grid"
    for i, text in enumerate(["Etat", "Ce que cela prouve", "Ecriture attendue", "Piece attendue"]):
        set_cell_text(proof.rows[0].cells[i], text, bold=True)
    for row in [
        ("Placement actif", "Le capital travaille encore.", "Placement PMG/FCP.", "Contrat ou mandat."),
        ("Liquidite interets", "Les interets sont exigibles mais non payes.", "liquidite_interets.", "Calcul et releve."),
        ("Liquidite capital", "Le capital est sorti du mandat mais non rembourse.", "liquidite_capital.", "Echeance ou decision de sortie."),
        ("Paiement client", "Le client a effectivement ete paye.", "payment/remboursement avec reference.", "Justificatif bancaire ou preuve de virement."),
    ]:
        cells = proof.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=8.5)
    style_table(proof)
    add_bullets(
        doc,
        [
            "A l'echeance PMG, le capital et les interets doivent passer en liquidite si le paiement n'est pas immediat.",
            "Le paiement des interets seul diminue la liquidite interets.",
            "Le paiement capital + interets diminue les deux liquidites concernees.",
            "La preuve de paiement doit inclure la date, la methode, la reference et le justificatif.",
            "Sans ecriture de paiement prouvee, la plateforme doit considerer que le client n'a pas encore ete paye.",
        ],
    )

    add_section(
        doc,
        "15. Compliance, Backoffice, DG, Admin et CRM",
        "Ces espaces servent a separer l'exploitation, le controle, la validation et le pilotage. Cette separation est importante pour la securite et l'audit.",
    )
    governance = doc.add_table(rows=1, cols=3)
    governance.style = "Table Grid"
    for i, text in enumerate(["Espace", "Actions principales", "Point d'attention"]):
        set_cell_text(governance.rows[0].cells[i], text, bold=True)
    for row in [
        ("Compliance", "Consulter clients, historiques, audits portefeuille, VL et releves envoyes.", "Ne modifie pas les donnees metier sans piste d'audit; sert de controle independant."),
        ("Backoffice", "Consulter et valider les transactions.", "Chaque validation doit etre rattachee a un utilisateur et une date."),
        ("Direction Generale", "Suivre le pilotage et les historiques de releves.", "Vue de supervision, pas de saisie operationnelle courante."),
        ("Admin Frontend", "Gerer utilisateurs, roles, menus et journaux.", "Limiter les permissions au strict besoin et surveiller les logs."),
        ("CRM", "Gerer prospects, statuts et clients commerciaux.", "Ne pas confondre prospect commercial et client financier actif."),
    ]:
        cells = governance.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=8.5)
    style_table(governance, LIGHT_BLUE)
    note(
        doc,
        "La securite parfaite n'existe pas en pratique; l'objectif est une securite robuste, auditable et testee. Les preuves de paiement, les permissions minimales, les journaux d'activite et les validations separees reduisent fortement le risque.",
        fill=LIGHT_GOLD,
        title="Securite.",
    )

    add_section(
        doc,
        "16. Controles mensuels recommandes",
        "Avant chaque campagne de releves, l'equipe doit effectuer un controle minimum pour eviter les ecarts de valorisation, les releves inutiles et les paiements non prouves.",
    )
    add_bullets(
        doc,
        [
            "Verifier la VL FCP du mois et le calcul au centime.",
            "Verifier les PMG a echeance dans le mois et ceux en periode de visibilite du 1er au 5.",
            "Controler les anniversaires PMG et les interets mensuels exceptionnels a mettre en liquidite.",
            "Comparer clients actifs, clients avec releve a envoyer et clients avec montant nul.",
            "Verifier que toute liquidite payee possede une preuve de paiement.",
            "Archiver ou conserver les releves envoyes pour audit.",
        ],
    )

    add_section(
        doc,
        "17. Depannage courant",
        "Les situations ci-dessous sont les anomalies les plus probables lors de l'exploitation courante.",
    )
    troubleshooting = doc.add_table(rows=1, cols=3)
    troubleshooting.style = "Table Grid"
    for i, text in enumerate(["Situation", "Cause probable", "Action"]):
        set_cell_text(troubleshooting.rows[0].cells[i], text, bold=True)
    for row in [
        ("Client affiche a zero", "Pas de placement actif dans le filtre ou placement expire/non valorise.", "Ouvrir Details puis Historique des transactions."),
        ("Nombre actifs different des releves", "La liste des releves utilise une logique de periode, produit et visibilite.", "Comparer filtre PMG/FCP, echeances et periode d'envoi."),
        ("Liquidite absente", "Ecriture liquidite_interets ou liquidite_capital non creee.", "Generer ou enregistrer le mouvement de liquidite attendu."),
        ("Paiement non prouve", "Absence de reference, methode ou justificatif.", "Exiger la preuve avant de considerer le client paye."),
        ("Ecart FCP au centime", "Arrondi trop tot sur VL ou nombre de parts.", "Recalculer avec toutes les decimales puis arrondir seulement a l'affichage."),
    ]:
        cells = troubleshooting.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=8.5)
    style_table(troubleshooting, LIGHT_GRAY)

    add_section(doc, "18. Glossaire")
    glossary = doc.add_table(rows=1, cols=2)
    glossary.style = "Table Grid"
    for i, text in enumerate(["Terme", "Definition"]):
        set_cell_text(glossary.rows[0].cells[i], text, bold=True)
    for row in [
        ("PMG", "Mandat ou produit de gestion avec capital, taux, date de debut, date de fin et modalite d'interets."),
        ("FCP", "Produit valorise par nombre de parts et valeur liquidative."),
        ("VL", "Valeur liquidative utilisee pour valoriser une part FCP."),
        ("Liquidite", "Montant disponible ou exigible, pas encore paye au client."),
        ("Precompte", "Interets constates des le debut selon la modalite du produit."),
        ("Releve", "Document envoye au client recapitulant la valorisation et les mouvements."),
    ]:
        cells = glossary.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, size=9)
        set_cell_text(cells[1], row[1], size=9)
    style_table(glossary)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
